import unittest
from io import BytesIO

from docx import Document

from app.services.cv_template_service import (
    build_replacement_map,
    find_unreplaced_placeholders,
    render_docx_template,
    safe_filename,
)


def _docx_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class CvTemplateServiceTests(unittest.TestCase):
    def test_render_docx_template_replaces_paragraph_table_and_header_tokens(self) -> None:
        template = Document()
        template.sections[0].header.paragraphs[0].text = "Code: {{surveyor_code}}"
        template.add_paragraph("Name: {{surveyor_name}}")
        table = template.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "{{phone_number}}"
        table.cell(0, 1).text = "{{assignments}}"

        replacements = build_replacement_map(
            {
                "surveyor_id": 7,
                "surveyor_code": "PPC-KBL-001",
                "surveyor_name": "Ahmad Wali",
                "phone_number": "+93700123456",
            },
            assignments=[
                {
                    "project_code": "P-001",
                    "project_name": "Survey Alpha",
                    "role": "LEAD",
                    "work_province_name": "Kabul",
                    "assignment_status": "ACTIVE",
                    "assignment_start_date": "2026-01-01",
                    "assignment_end_date": None,
                    "is_current_active": True,
                }
            ],
        )

        rendered_bytes = render_docx_template(_docx_bytes(template), replacements)
        rendered = Document(BytesIO(rendered_bytes))

        self.assertTrue(rendered_bytes.startswith(b"PK"))
        self.assertIn("PPC-KBL-001", rendered.sections[0].header.paragraphs[0].text)
        self.assertEqual(rendered.paragraphs[0].text, "Name: Ahmad Wali")
        self.assertEqual(rendered.tables[0].cell(0, 0).text, "+93700123456")
        self.assertIn("Survey Alpha", rendered.tables[0].cell(0, 1).text)
        self.assertEqual(find_unreplaced_placeholders(rendered_bytes), [])

    def test_find_unreplaced_placeholders_returns_unknown_tokens(self) -> None:
        template = Document()
        template.add_paragraph("Custom: {{unknown_field}}")
        rendered_bytes = render_docx_template(_docx_bytes(template), {})

        self.assertEqual(find_unreplaced_placeholders(rendered_bytes), ["{{unknown_field}}"])

    def test_render_docx_template_smart_fills_standard_form_h_tables(self) -> None:
        template = Document()
        template.add_paragraph("FORM H: FORMAT FOR CV OF PROPOSED KEY PERSONNEL")
        personnel_table = template.add_table(rows=11, cols=3)
        personnel_table.cell(0, 0).text = "Position (as per ToR)"
        personnel_table.cell(1, 1).text = "Name:"
        personnel_table.cell(2, 1).text = "Nationality:"
        personnel_table.cell(6, 1).text = "Telephone:"
        personnel_table.cell(6, 2).text = "Email:"
        personnel_table.cell(7, 1).text = "Job title:"
        experience_table = template.add_table(rows=2, cols=3)
        experience_table.cell(0, 0).text = "From"
        experience_table.cell(0, 1).text = "To"
        experience_table.cell(0, 2).text = "Company / Project / Position / Relevant technical and management experience"

        assignments = [
            {
                "project_code": "P-001",
                "project_name": "Survey Alpha",
                "role": "LEAD_SURVEYOR",
                "work_province_name": "Kabul",
                "assignment_status": "ACTIVE",
                "assignment_start_date": "2026-01-01",
                "assignment_end_date": None,
                "is_current_active": True,
            }
        ]
        replacements = build_replacement_map(
            {
                "surveyor_id": 7,
                "surveyor_code": "PPC-KBL-001",
                "surveyor_name": "Ahmad Wali",
                "phone_number": "+93700123456",
                "email_address": "ahmad@example.com",
                "nationality": "Afghan",
            },
            assignments=assignments,
        )

        rendered_bytes = render_docx_template(_docx_bytes(template), replacements, assignments=assignments)
        rendered = Document(BytesIO(rendered_bytes))

        self.assertIn("Lead Surveyor", rendered.tables[0].cell(0, 1).text)
        self.assertEqual(rendered.tables[0].cell(1, 1).text, "Name: Ahmad Wali")
        self.assertEqual(rendered.tables[0].cell(2, 1).text, "Nationality: Afghan")
        self.assertEqual(rendered.tables[0].cell(6, 1).text, "Telephone: +93700123456")
        self.assertEqual(rendered.tables[0].cell(6, 2).text, "Email: ahmad@example.com")
        self.assertEqual(rendered.tables[1].cell(1, 0).text, "2026-01-01")
        self.assertEqual(rendered.tables[1].cell(1, 1).text, "Present")
        self.assertIn("Survey Alpha", rendered.tables[1].cell(1, 2).text)

    def test_safe_filename_removes_unsafe_characters(self) -> None:
        self.assertEqual(safe_filename("PPC/Kabul: 001"), "PPC_Kabul_001")


if __name__ == "__main__":
    unittest.main()
