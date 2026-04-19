from __future__ import annotations

import re
import shutil
import subprocess
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from docx import Document

PLACEHOLDER_FIELDS = [
    ("surveyor_name", "Surveyor full name"),
    ("surveyor_code", "Surveyor code"),
    ("gender", "Gender"),
    ("father_name", "Father name"),
    ("tazkira_no", "Tazkira number"),
    ("email_address", "Email address"),
    ("phone_number", "Phone number"),
    ("whatsapp_number", "WhatsApp number"),
    ("permanent_province_name", "Permanent province"),
    ("current_province_name", "Current province"),
    ("cv_link", "Existing CV link"),
    ("default_bank_name", "Default bank"),
    ("default_payment_type", "Default payment type"),
    ("default_payout_value", "Default payout account or mobile"),
    ("bank_accounts", "All linked bank accounts"),
    ("assignments", "All surveyor project assignments"),
    ("active_assignments", "Only active/current project assignments"),
    ("project_names", "Project names list"),
    ("current_role", "Current or latest assignment role"),
    ("current_project_name", "Current or latest project name"),
    ("current_employer", "Current project client or partner"),
    ("proposer_name", "Form H proposer name"),
    ("rfp_reference", "Form H RFP reference"),
    ("position_tor", "Position as per ToR"),
    ("nationality", "Nationality"),
    ("date_of_birth", "Date of birth"),
    ("language_proficiency", "Language proficiency"),
    ("education_qualifications", "Education and qualifications"),
    ("professional_certifications", "Professional certifications"),
    ("references", "References"),
    ("generated_date", "Document generated date"),
]

WORD_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_TEMPLATE_SIZE_BYTES = 8 * 1024 * 1024
MAX_TEMPLATE_SIZE_MB = MAX_TEMPLATE_SIZE_BYTES // (1024 * 1024)


def _display(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    text = str(value).strip()
    return text if text else fallback


def _format_status(value: Any) -> str:
    text = _display(value)
    return text.replace("_", " ").title() if text else ""


def _format_date_range(start_value: Any, end_value: Any) -> str:
    start_text = _display(start_value, "No start date")
    end_text = _display(end_value, "Present")
    return f"{start_text} - {end_text}"


def _format_bank_account(account: dict[str, Any]) -> str:
    bank_name = _display(account.get("bank_name"), "No bank")
    payment_type = _format_status(account.get("payment_type")) or "Payment"
    payout_value = _display(account.get("account_number")) or _display(account.get("mobile_number"), "No payout value")
    account_title = _display(account.get("account_title"))
    status = "Active" if account.get("is_active") else "Inactive"
    default_marker = "Default" if account.get("is_default") else "Secondary"
    title_part = f" - {account_title}" if account_title else ""
    return f"{bank_name} - {payment_type} - {payout_value}{title_part} - {status} - {default_marker}"


def _format_assignment(assignment: dict[str, Any]) -> str:
    project_code = _display(assignment.get("project_code"))
    project_name = _display(assignment.get("project_name"), "Unnamed project")
    role = _format_status(assignment.get("role")) or "Surveyor"
    status = _format_status(assignment.get("assignment_status")) or "Unspecified"
    province = _display(assignment.get("work_province_name")) or _display(assignment.get("work_province_code"), "No province")
    date_range = _format_date_range(assignment.get("assignment_start_date"), assignment.get("assignment_end_date"))
    project_label = f"{project_code} - {project_name}" if project_code else project_name
    return f"{project_label} | {role} | {province} | {status} | {date_range}"


def _join_lines(values: list[str], fallback: str = "") -> str:
    cleaned = [value for value in values if value]
    return "\n".join(cleaned) if cleaned else fallback


def _current_assignment(assignments: list[dict[str, Any]]) -> dict[str, Any]:
    active_assignments = [assignment for assignment in assignments if assignment.get("is_current_active")]
    if active_assignments:
        return active_assignments[0]
    if assignments:
        return assignments[0]
    return {}


def _years_since(value: Any) -> str:
    if not isinstance(value, date):
        return ""
    today = date.today()
    years = today.year - value.year - ((today.month, today.day) < (value.month, value.day))
    return str(max(years, 0))


def build_replacement_map(
    profile: dict[str, Any],
    bank_accounts: list[dict[str, Any]] | None = None,
    assignments: list[dict[str, Any]] | None = None,
) -> dict[str, str]:
    bank_accounts = bank_accounts or []
    assignments = assignments or []
    active_assignments = [assignment for assignment in assignments if assignment.get("is_current_active")]
    current_assignment = _current_assignment(assignments)
    current_role = _format_status(current_assignment.get("role"))
    current_employer = _display(current_assignment.get("implementing_partner")) or _display(current_assignment.get("client_name"))

    replacements = {
        "surveyor_id": _display(profile.get("surveyor_id")),
        "surveyor_code": _display(profile.get("surveyor_code")),
        "surveyor_name": _display(profile.get("surveyor_name")),
        "gender": _format_status(profile.get("gender")),
        "father_name": _display(profile.get("father_name")),
        "tazkira_no": _display(profile.get("tazkira_no")),
        "email_address": _display(profile.get("email_address")),
        "whatsapp_number": _display(profile.get("whatsapp_number")),
        "phone_number": _display(profile.get("phone_number")),
        "permanent_province_code": _display(profile.get("permanent_province_code")),
        "permanent_province_name": _display(profile.get("permanent_province_name")),
        "current_province_code": _display(profile.get("current_province_code")),
        "current_province_name": _display(profile.get("current_province_name")),
        "cv_link": _display(profile.get("cv_link")),
        "cv_file_name": _display(profile.get("cv_file_name")),
        "tazkira_image_name": _display(profile.get("tazkira_image_name")),
        "tazkira_pdf_name": _display(profile.get("tazkira_pdf_name")),
        "tazkira_word_name": _display(profile.get("tazkira_word_name")),
        "document_count": _display(profile.get("document_count"), "0"),
        "account_count": _display(profile.get("account_count"), "0"),
        "active_account_count": _display(profile.get("active_account_count"), "0"),
        "default_bank_name": _display(profile.get("default_bank_name")),
        "default_payment_type": _format_status(profile.get("default_payment_type")),
        "default_payout_value": _display(profile.get("default_payout_value")),
        "bank_names": _display(profile.get("bank_names")),
        "bank_accounts": _join_lines([_format_bank_account(account) for account in bank_accounts], "No linked bank accounts"),
        "assignments": _join_lines([_format_assignment(assignment) for assignment in assignments], "No project assignments"),
        "active_assignments": _join_lines(
            [_format_assignment(assignment) for assignment in active_assignments],
            "No active project assignments",
        ),
        "project_names": _join_lines([_display(assignment.get("project_name")) for assignment in assignments], "No project assignments"),
        "current_role": current_role,
        "current_project_name": _display(current_assignment.get("project_name")),
        "current_employer": current_employer,
        "current_employer_address": _display(current_assignment.get("work_province_name")) or _display(profile.get("current_province_name")),
        "years_with_present_employer": _years_since(current_assignment.get("assignment_start_date")),
        "position_tor": current_role,
        "proposer_name": _display(profile.get("proposer_name")),
        "rfp_reference": _display(profile.get("rfp_reference")),
        "nationality": _display(profile.get("nationality")),
        "date_of_birth": _display(profile.get("date_of_birth")),
        "language_proficiency": _display(profile.get("language_proficiency")),
        "education_qualifications": _display(profile.get("education_qualifications")),
        "professional_certifications": _display(profile.get("professional_certifications")),
        "references": _display(profile.get("references")),
        "generated_date": date.today().strftime("%Y-%m-%d"),
    }
    return replacements


def replacement_tokens(replacements: dict[str, str]) -> dict[str, str]:
    return {f"{{{{{key}}}}}": value for key, value in replacements.items()}


def _replace_in_paragraph(paragraph, tokens: dict[str, str]) -> None:
    original_text = paragraph.text
    if "{{" not in original_text:
        return

    updated_text = original_text
    for token, value in tokens.items():
        updated_text = updated_text.replace(token, value)

    if updated_text == original_text:
        return

    if not paragraph.runs:
        paragraph.add_run(updated_text)
        return

    paragraph.runs[0].text = updated_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_table(table, tokens: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, tokens)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, tokens)


def _replace_in_document(document, tokens: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, tokens)
    for table in document.tables:
        _replace_in_table(table, tokens)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _replace_in_paragraph(paragraph, tokens)
            for table in part.tables:
                _replace_in_table(table, tokens)


def _set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        _set_paragraph_text(cell.paragraphs[0], text)
        for paragraph in cell.paragraphs[1:]:
            _set_paragraph_text(paragraph, "")
    else:
        cell.text = text


def _cell_text(cell) -> str:
    return (cell.text or "").strip()


def _is_form_h_document(document) -> bool:
    for paragraph in document.paragraphs:
        if "FORMAT FOR CV OF PROPOSED KEY PERSONNEL" in paragraph.text.upper():
            return True
    return any(
        table.rows
        and table.columns
        and "POSITION (AS PER TOR)" in _cell_text(table.cell(0, 0)).upper()
        for table in document.tables
    )


def _field(label: str, value: str) -> str:
    return f"{label}: {value}" if value else f"{label}:"


def _fill_value_cell(table, row_index: int, column_index: int, label: str, value: str) -> None:
    if len(table.rows) <= row_index or len(table.columns) <= column_index:
        return
    if value:
        _set_cell_text(table.cell(row_index, column_index), _field(label, value))


def _fill_header_table(table, replacements: dict[str, str]) -> None:
    if len(table.rows) < 2 or len(table.columns) < 4:
        return
    proposer_name = replacements.get("proposer_name", "")
    rfp_reference = replacements.get("rfp_reference", "")
    generated_date = replacements.get("generated_date", "")

    if "NAME OF PROPOSER" in _cell_text(table.cell(0, 0)).upper() and proposer_name:
        _set_cell_text(table.cell(0, 1), proposer_name)
    if "DATE" in _cell_text(table.cell(0, 2)).upper() and generated_date:
        _set_cell_text(table.cell(0, 3), generated_date)
    if "RFP REFERENCE" in _cell_text(table.cell(1, 0)).upper() and rfp_reference:
        _set_cell_text(table.cell(1, 1), rfp_reference)


def _fill_personnel_table(table, replacements: dict[str, str]) -> None:
    if len(table.rows) < 11 or len(table.columns) < 3:
        return
    if "POSITION (AS PER TOR)" not in _cell_text(table.cell(0, 0)).upper():
        return

    position = replacements.get("position_tor") or replacements.get("current_role", "")
    if position:
        _set_cell_text(table.cell(0, 1), position)

    _fill_value_cell(table, 1, 1, "Name", replacements.get("surveyor_name", ""))
    _fill_value_cell(table, 2, 1, "Nationality", replacements.get("nationality", ""))
    _fill_value_cell(table, 2, 2, "Date of birth", replacements.get("date_of_birth", ""))
    _fill_value_cell(table, 3, 1, "Language Proficiency", replacements.get("language_proficiency", ""))
    _fill_value_cell(table, 4, 1, "Name of employer", replacements.get("current_employer", ""))
    _fill_value_cell(table, 4, 2, "Contact", replacements.get("whatsapp_number", "") or replacements.get("phone_number", ""))
    _fill_value_cell(table, 5, 1, "Address of employer", replacements.get("current_employer_address", ""))
    _fill_value_cell(table, 6, 1, "Telephone", replacements.get("phone_number", ""))
    _fill_value_cell(table, 6, 2, "Email", replacements.get("email_address", ""))
    _fill_value_cell(table, 7, 1, "Job title", replacements.get("current_role", ""))
    _fill_value_cell(table, 7, 2, "Years with present employer", replacements.get("years_with_present_employer", ""))

    if replacements.get("education_qualifications"):
        _set_cell_text(table.cell(8, 1), replacements["education_qualifications"])
    if replacements.get("professional_certifications"):
        _set_cell_text(table.cell(9, 1), replacements["professional_certifications"])
    if replacements.get("references"):
        _set_cell_text(table.cell(10, 1), replacements["references"])


def _assignment_detail(assignment: dict[str, Any]) -> str:
    project_name = _display(assignment.get("project_name"), "Unnamed project")
    project_code = _display(assignment.get("project_code"))
    role = _format_status(assignment.get("role")) or "Surveyor"
    province = _display(assignment.get("work_province_name")) or _display(assignment.get("work_province_code"))
    client = _display(assignment.get("client_name"))
    partner = _display(assignment.get("implementing_partner"))
    status = _format_status(assignment.get("assignment_status"))
    parts = []
    if project_code:
        parts.append(project_code)
    parts.append(project_name)
    parts.append(f"Position: {role}")
    if province:
        parts.append(f"Province: {province}")
    if client:
        parts.append(f"Client: {client}")
    if partner:
        parts.append(f"Partner: {partner}")
    if status:
        parts.append(f"Status: {status}")
    return " | ".join(parts)


def _fill_experience_table(table, assignments: list[dict[str, Any]]) -> None:
    if len(table.rows) < 2 or len(table.columns) < 3:
        return
    headers = [_cell_text(table.cell(0, column_index)).upper() for column_index in range(3)]
    if headers[:2] != ["FROM", "TO"] or "COMPANY" not in headers[2]:
        return

    selected_assignments = assignments[:20]
    if not selected_assignments:
        _set_cell_text(table.cell(1, 0), "")
        _set_cell_text(table.cell(1, 1), "")
        _set_cell_text(table.cell(1, 2), "")
        return

    while len(table.rows) < len(selected_assignments) + 1:
        table.add_row()

    for offset, assignment in enumerate(selected_assignments, start=1):
        _set_cell_text(table.cell(offset, 0), _display(assignment.get("assignment_start_date")))
        _set_cell_text(table.cell(offset, 1), _display(assignment.get("assignment_end_date"), "Present"))
        _set_cell_text(table.cell(offset, 2), _assignment_detail(assignment))


def _fill_availability_table(table, assignments: list[dict[str, Any]]) -> None:
    if len(table.rows) < 2 or len(table.columns) < 2:
        return
    headers = [_cell_text(table.cell(0, column_index)).upper() for column_index in range(2)]
    if headers != ["FROM", "TO"]:
        return

    active_assignments = [assignment for assignment in assignments if assignment.get("is_current_active")]
    selected_assignments = (active_assignments or assignments)[:10]
    if not selected_assignments:
        for row_index in range(1, len(table.rows)):
            _set_cell_text(table.cell(row_index, 0), "")
            _set_cell_text(table.cell(row_index, 1), "")
        return

    while len(table.rows) < len(selected_assignments) + 1:
        table.add_row()

    for offset, assignment in enumerate(selected_assignments, start=1):
        _set_cell_text(table.cell(offset, 0), _display(assignment.get("assignment_start_date")))
        _set_cell_text(table.cell(offset, 1), _display(assignment.get("assignment_end_date"), "Present"))

    for row_index in range(len(selected_assignments) + 1, len(table.rows)):
        _set_cell_text(table.cell(row_index, 0), "")
        _set_cell_text(table.cell(row_index, 1), "")


def _fill_signature_paragraphs(document, replacements: dict[str, str]) -> None:
    name = replacements.get("surveyor_name", "")
    title = replacements.get("current_role") or replacements.get("position_tor", "")
    generated_date = replacements.get("generated_date", "")
    for paragraph in document.paragraphs:
        normalized = " ".join(paragraph.text.split()).upper()
        if normalized == "NAME: TITLE: DATE: SIGNATURE:" and name:
            _set_paragraph_text(paragraph, f"Name: {name}\tTitle: {title}\tDate: {generated_date}\tSignature:")


def _auto_fill_standard_cv_forms(
    document,
    replacements: dict[str, str],
    assignments: list[dict[str, Any]],
) -> None:
    if not _is_form_h_document(document):
        return

    for table in document.tables:
        first_cell = _cell_text(table.cell(0, 0)).upper() if table.rows and table.columns else ""
        if "NAME OF PROPOSER" in first_cell:
            _fill_header_table(table, replacements)
        elif "POSITION (AS PER TOR)" in first_cell:
            _fill_personnel_table(table, replacements)
        elif first_cell == "FROM":
            if len(table.columns) >= 3:
                _fill_experience_table(table, assignments)
            else:
                _fill_availability_table(table, assignments)

    _fill_signature_paragraphs(document, replacements)


def render_docx_template(
    template_bytes: bytes,
    replacements: dict[str, str],
    *,
    assignments: list[dict[str, Any]] | None = None,
    smart_fill: bool = True,
) -> bytes:
    document = Document(BytesIO(template_bytes))
    _replace_in_document(document, replacement_tokens(replacements))
    if smart_fill:
        _auto_fill_standard_cv_forms(document, replacements, assignments or [])
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def find_unreplaced_placeholders(document_bytes: bytes) -> list[str]:
    document = Document(BytesIO(document_bytes))
    found: set[str] = set()
    pattern = re.compile(r"\{\{[a-zA-Z0-9_]+\}\}")

    def collect_text(text: str) -> None:
        for match in pattern.findall(text or ""):
            found.add(match)

    for paragraph in document.paragraphs:
        collect_text(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_text(paragraph.text)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                collect_text(paragraph.text)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            collect_text(paragraph.text)

    return sorted(found)


def available_pdf_converter() -> str | None:
    return shutil.which("soffice") or shutil.which("libreoffice")


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes | None:
    converter = available_pdf_converter()
    if not converter:
        return None

    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        source_path = temp_path / "generated_cv.docx"
        pdf_path = temp_path / "generated_cv.pdf"
        source_path.write_bytes(docx_bytes)
        subprocess.run(
            [converter, "--headless", "--convert-to", "pdf", "--outdir", str(temp_path), str(source_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=90,
        )
        if not pdf_path.exists():
            return None
        return pdf_path.read_bytes()


def safe_filename(value: str, fallback: str = "surveyor_cv") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value or "").strip("._-")
    return cleaned or fallback
