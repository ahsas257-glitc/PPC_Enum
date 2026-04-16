from datetime import date, datetime
from io import BytesIO

from docx import Document
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
import streamlit as st

from app.core.constants import ASSIGNMENT_STATUSES, PROJECT_STATUSES, PROJECT_TYPES
from app.core.permissions import ensure_role
from app.core.session import get_current_user
from app.design.components.cards import render_hero, render_panel_intro, render_stat_band
from app.design.components.filters import apply_text_filter
from app.design.components.tables import render_table
from app.design.components import validation as vf
from app.repositories.province_repository import ProvinceRepository
from app.services.project_service import ProjectService
from app.services.surveyor_service import SurveyorService

PROJECT_FORM_KEY = "project_form"
PROJECT_NAME_KEY = "project_form_name"
PROJECT_SHORT_NAME_KEY = "project_form_short_name"
PROJECT_CLIENT_KEY = "project_form_client"
PROJECT_PARTNER_KEY = "project_form_partner"
PROJECT_TYPE_KEY = "project_form_type"
PROJECT_STATUS_KEY = "project_form_status"
PROJECT_START_KEY = "project_form_start"
PROJECT_END_KEY = "project_form_end"
PROJECT_DOCUMENT_KEY = "project_form_document"
PROJECT_NOTES_KEY = "project_form_notes"
PROJECT_OVERVIEW_STATUS_FILTER_KEY = "project_overview_status_filter"
PROJECT_OVERVIEW_CLIENT_FILTER_KEY = "project_overview_client_filter"
PROJECT_OVERVIEW_TYPE_FILTER_KEY = "project_overview_type_filter"

ASSIGNMENT_FORM_KEY = "assignment_form"
ASSIGNMENT_PROJECT_KEY = "assignment_form_project"
ASSIGNMENT_SURVEYORS_KEY = "assignment_form_surveyors"
ASSIGNMENT_ROLE_KEY = "assignment_form_role"
ASSIGNMENT_STATUS_KEY = "assignment_form_status"
ASSIGNMENT_PROVINCE_KEY = "assignment_form_province"
ASSIGNMENT_PROVINCES_KEY = "assignment_form_provinces"
ASSIGNMENT_EXTRA_PROVINCES_KEY = "assignment_form_extra_provinces"
ASSIGNMENT_START_KEY = "assignment_form_start"
ASSIGNMENT_END_KEY = "assignment_form_end"
ASSIGNMENT_NOTES_KEY = "assignment_form_notes"


def _candidate_selection_label(item: dict) -> str:
    return (
        f"{item['surveyor_name']} ({item['surveyor_code']})  "
        f"Docs {item.get('document_count', 0)}/4  "
        f"Accounts {item.get('active_account_count', 0)}  "
        f"Projects {item.get('active_project_count', 0)}"
    )


def _selection_control_key(prefix: str, value: object) -> str:
    return f"{prefix}_{value}"


def _render_selection_spacer(height: str = "0.58rem") -> None:
    st.html(f'<div style="height:{height};"></div>')


def _render_checkbox_selector(
    title: str,
    options: list[tuple[str, str]],
    *,
    state_key: str,
    control_prefix: str,
    help_text: str | None = None,
    expanded: bool = True,
    columns: int = 2,
) -> list[str]:
    valid_values = {value for value, _ in options}
    selected_values = [value for value in st.session_state.get(state_key, []) if value in valid_values]

    for value, _ in options:
        control_key = _selection_control_key(control_prefix, value)
        if control_key not in st.session_state:
            st.session_state[control_key] = value in selected_values

    with st.expander(title, expanded=expanded):
        if help_text:
            st.caption(help_text)
        if not options:
            st.caption("Nothing is available here yet.")
        else:
            selector_columns = st.columns(columns)
            for index, (value, label) in enumerate(options):
                with selector_columns[index % columns]:
                    st.checkbox(label, key=_selection_control_key(control_prefix, value))

    selected_values = [
        value
        for value, _ in options
        if st.session_state.get(_selection_control_key(control_prefix, value), False)
    ]
    st.session_state[state_key] = selected_values
    return selected_values


def _render_checkbox_grid(
    options: list[tuple[str, str]],
    *,
    control_prefix: str,
    selected_values: list[str] | None = None,
    columns: int = 2,
) -> list[str]:
    valid_values = {value for value, _ in options}
    selected_lookup = {value for value in (selected_values or []) if value in valid_values}

    for value, _ in options:
        control_key = _selection_control_key(control_prefix, value)
        if control_key not in st.session_state:
            st.session_state[control_key] = value in selected_lookup

    selector_columns = st.columns(columns)
    for index, (value, label) in enumerate(options):
        with selector_columns[index % columns]:
            st.checkbox(label, key=_selection_control_key(control_prefix, value))

    return [
        value
        for value, _ in options
        if st.session_state.get(_selection_control_key(control_prefix, value), False)
    ]


def _set_checkbox_grid_values(options: list[tuple[str, str]], *, control_prefix: str, checked: bool) -> None:
    for value, _ in options:
        st.session_state[_selection_control_key(control_prefix, value)] = checked


def _read_checkbox_grid_values(options: list[tuple[str, str]], *, control_prefix: str) -> list[str]:
    return [
        value
        for value, _ in options
        if st.session_state.get(_selection_control_key(control_prefix, value), False)
    ]


def _filter_assignment_rows(rows: list[dict], query_text: str) -> list[dict]:
    term = (query_text or "").strip().lower()
    if not term:
        return rows
    filtered_rows: list[dict] = []
    for item in rows:
        haystack = " ".join(
            [
                str(item.get("surveyor_name") or ""),
                str(item.get("surveyor_code") or ""),
                str(item.get("availability_province_name") or ""),
                str(item.get("availability_province_code") or ""),
            ]
        ).lower()
        if term in haystack:
            filtered_rows.append(item)
    return filtered_rows


def _province_summary_frame(candidate_rows: list[dict]) -> pd.DataFrame:
    if not candidate_rows:
        return pd.DataFrame()
    frame = pd.DataFrame(candidate_rows)
    summary = (
        frame.groupby("availability_province_name", dropna=False)
        .agg(
            surveyor_count=("surveyor_id", "count"),
            complete_docs=("document_count", lambda values: int((pd.Series(values) >= 4).sum())),
            active_accounts=("active_account_count", "sum"),
            active_projects=("active_project_count", "sum"),
        )
        .reset_index()
        .rename(
            columns={
                "availability_province_name": "Province",
                "surveyor_count": "Surveyors",
                "complete_docs": "Ready docs",
                "active_accounts": "Active accounts",
                "active_projects": "Active projects",
            }
        )
    )
    return summary.sort_values(["Surveyors", "Province"], ascending=[False, True], ignore_index=True)


def _group_assignment_candidates(candidate_rows: list[dict]) -> list[tuple[str, str, list[dict]]]:
    grouped: dict[tuple[str, str], list[dict]] = {}
    for item in candidate_rows:
        province_code = item.get("availability_province_code") or "NA"
        province_name = item.get("availability_province_name") or "Province not set"
        grouped.setdefault((province_code, province_name), []).append(item)

    ordered_groups: list[tuple[str, str, list[dict]]] = []
    for (province_code, province_name), rows in sorted(grouped.items(), key=lambda group: (group[0][1], group[0][0])):
        ordered_groups.append(
            (
                province_code,
                province_name,
                sorted(rows, key=lambda item: (str(item.get("surveyor_name") or "").lower(), -int(item.get("surveyor_id") or 0))),
            )
        )
    return ordered_groups


def _assignment_conflict_frame(conflicts: list[dict]) -> pd.DataFrame:
    if not conflicts:
        return pd.DataFrame()
    frame = pd.DataFrame(conflicts)[
        [
            "surveyor_code",
            "surveyor_name",
            "project_code",
            "project_name",
            "work_province_name",
            "assignment_status",
            "assignment_start_date",
            "assignment_end_date",
            "same_project",
            "overlaps_window",
        ]
    ].rename(
        columns={
            "surveyor_code": "Surveyor code",
            "surveyor_name": "Surveyor",
            "project_code": "Project code",
            "project_name": "Project",
            "work_province_name": "Province",
            "assignment_status": "Status",
            "assignment_start_date": "Start",
            "assignment_end_date": "End",
            "same_project": "Same project",
            "overlaps_window": "Overlaps dates",
        }
    )
    return frame


def _project_status_overview_frame(projects_df: pd.DataFrame) -> pd.DataFrame:
    if projects_df.empty:
        return pd.DataFrame()
    frame = projects_df.copy()
    frame["assignment_count"] = pd.to_numeric(frame.get("assignment_count"), errors="coerce").fillna(0).astype(int)
    overview = (
        frame.groupby("status", dropna=False)
        .agg(
            project_count=("project_id", "count"),
            assignment_count=("assignment_count", "sum"),
            client_count=("client_name", lambda values: int(pd.Series(values).fillna("").replace("", pd.NA).dropna().nunique())),
        )
        .reset_index()
        .rename(
            columns={
                "status": "Status",
                "project_count": "Projects",
                "assignment_count": "Assignments",
                "client_count": "Clients",
            }
        )
    )
    return overview.sort_values("Projects", ascending=False, ignore_index=True)


def _project_type_overview_frame(projects_df: pd.DataFrame) -> pd.DataFrame:
    if projects_df.empty:
        return pd.DataFrame()
    frame = projects_df.copy()
    frame["assignment_count"] = pd.to_numeric(frame.get("assignment_count"), errors="coerce").fillna(0).astype(int)
    overview = (
        frame.groupby("project_type", dropna=False)
        .agg(
            project_count=("project_id", "count"),
            active_projects=("status", lambda values: int((pd.Series(values) == "ACTIVE").sum())),
            assignment_count=("assignment_count", "sum"),
        )
        .reset_index()
        .rename(
            columns={
                "project_type": "Project type",
                "project_count": "Projects",
                "active_projects": "Active",
                "assignment_count": "Assignments",
            }
        )
    )
    return overview.sort_values("Projects", ascending=False, ignore_index=True)


def _project_client_overview_frame(projects_df: pd.DataFrame) -> pd.DataFrame:
    if projects_df.empty:
        return pd.DataFrame()
    frame = projects_df.copy()
    frame["client_name"] = frame["client_name"].fillna("Unspecified")
    frame["assignment_count"] = pd.to_numeric(frame.get("assignment_count"), errors="coerce").fillna(0).astype(int)
    overview = (
        frame.groupby("client_name", dropna=False)
        .agg(
            project_count=("project_id", "count"),
            active_projects=("status", lambda values: int((pd.Series(values) == "ACTIVE").sum())),
            planned_projects=("status", lambda values: int((pd.Series(values) == "PLANNED").sum())),
            assignment_count=("assignment_count", "sum"),
        )
        .reset_index()
        .rename(
            columns={
                "client_name": "Client",
                "project_count": "Projects",
                "active_projects": "Active",
                "planned_projects": "Planned",
                "assignment_count": "Assignments",
            }
        )
    )
    return overview.sort_values(["Projects", "Client"], ascending=[False, True], ignore_index=True)


def _project_stat_items(projects_df: pd.DataFrame) -> list[tuple[str, str]]:
    if projects_df.empty:
        return [
            ("Total projects", "0"),
            ("Active", "0"),
            ("Planned", "0"),
            ("On hold", "0"),
            ("Closed", "0"),
            ("Assignments", "0"),
        ]
    frame = projects_df.copy()
    statuses = frame["status"].fillna("")
    total_assignments = int(pd.to_numeric(frame.get("assignment_count"), errors="coerce").fillna(0).sum())
    return [
        ("Total projects", f"{len(frame)}"),
        ("Active", f"{int((statuses == 'ACTIVE').sum())}"),
        ("Planned", f"{int((statuses == 'PLANNED').sum())}"),
        ("On hold", f"{int((statuses == 'ON_HOLD').sum())}"),
        ("Closed", f"{int((statuses == 'CLOSED').sum())}"),
        ("Assignments", f"{total_assignments}"),
    ]


def _export_cell_value(value: object) -> object:
    if value is None or pd.isna(value):
        return ""
    if isinstance(value, pd.Timestamp):
        value = value.to_pydatetime()
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, (list, tuple, set)):
        return ", ".join(str(item) for item in value if item is not None)
    if isinstance(value, bool):
        return "Yes" if value else "No"
    return value


def _export_ready_frame(frame: pd.DataFrame) -> pd.DataFrame:
    if frame.empty:
        return frame.copy()
    export_frame = frame.copy()
    export_frame.columns = [str(column).replace("_", " ").strip().title() for column in export_frame.columns]
    for column_name in export_frame.columns:
        export_frame[column_name] = export_frame[column_name].map(_export_cell_value)
    return export_frame


def _filters_summary_text(statuses: list[str], clients: list[str], project_types: list[str]) -> str:
    parts = [
        f"Statuses: {', '.join(statuses) if statuses else 'All'}",
        f"Clients: {', '.join(clients) if clients else 'All'}",
        f"Project types: {', '.join(project_types) if project_types else 'All'}",
    ]
    return " | ".join(parts)


def _dataframe_to_csv_bytes(frame: pd.DataFrame) -> bytes:
    return _export_ready_frame(frame).to_csv(index=False).encode("utf-8-sig")


def _frames_to_xlsx_bytes(sheets: dict[str, pd.DataFrame]) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for sheet_name, frame in sheets.items():
            _export_ready_frame(frame).to_excel(writer, index=False, sheet_name=sheet_name[:31] or "Sheet1")
    return output.getvalue()


def _summary_table_data(summary_items: list[tuple[str, str]]) -> list[list[str]]:
    return [["Metric", "Value"]] + [[label, value] for label, value in summary_items]


def _build_word_report_bytes(
    title: str,
    subtitle: str,
    summary_items: list[tuple[str, str]],
    sections: list[tuple[str, pd.DataFrame]],
) -> bytes:
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(subtitle)

    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Metric"
    summary_table.rows[0].cells[1].text = "Value"
    for label, value in summary_items:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = str(label)
        row_cells[1].text = str(value)

    for section_title, frame in sections:
        document.add_heading(section_title, level=1)
        export_frame = _export_ready_frame(frame)
        if export_frame.empty:
            document.add_paragraph("No data available.")
            continue
        table = document.add_table(rows=1, cols=len(export_frame.columns))
        table.style = "Table Grid"
        for index, column_name in enumerate(export_frame.columns):
            table.rows[0].cells[index].text = str(column_name)
        for row in export_frame.itertuples(index=False):
            row_cells = table.add_row().cells
            for index, value in enumerate(row):
                row_cells[index].text = str(value)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _build_pdf_report_bytes(
    title: str,
    subtitle: str,
    summary_items: list[tuple[str, str]],
    sections: list[tuple[str, pd.DataFrame]],
) -> bytes:
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        leftMargin=24,
        rightMargin=24,
        topMargin=24,
        bottomMargin=24,
    )
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 8), Paragraph(subtitle, styles["BodyText"]), Spacer(1, 12)]

    summary_table = Table(_summary_table_data(summary_items), repeatRows=1)
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3b5b")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#a9b9cc")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.extend([summary_table, Spacer(1, 16)])

    for section_title, frame in sections:
        story.append(Paragraph(section_title, styles["Heading2"]))
        story.append(Spacer(1, 6))
        export_frame = _export_ready_frame(frame)
        if export_frame.empty:
            story.extend([Paragraph("No data available.", styles["BodyText"]), Spacer(1, 12)])
            continue
        table_data = [list(export_frame.columns)] + export_frame.astype(str).values.tolist()
        section_table = Table(table_data, repeatRows=1)
        section_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#234765")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#b6c4d2")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fb")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.extend([section_table, Spacer(1, 16)])

    document.build(story)
    return output.getvalue()


def _filter_projects_overview(
    projects_df: pd.DataFrame,
    *,
    statuses: list[str],
    clients: list[str],
    project_types: list[str],
) -> pd.DataFrame:
    if projects_df.empty:
        return projects_df
    filtered = projects_df.copy()
    if statuses:
        filtered = filtered[filtered["status"].isin(statuses)]
    if clients:
        filtered = filtered[filtered["client_name"].fillna("Unspecified").isin(clients)]
    if project_types:
        filtered = filtered[filtered["project_type"].isin(project_types)]
    return filtered.reset_index(drop=True)


def render_projects_page() -> None:
    ensure_role("super_admin", "admin", "manager")
    project_service = ProjectService()

    render_hero(
        "Operations",
        kicker="Projects",
    )
    active_view = st.radio(
        "Projects view",
        ["Project Overview", "New Project", "Assignments", "Project Data"],
        key="projects_active_view",
        horizontal=True,
        label_visibility="collapsed",
    )

    if active_view == "Project Overview":
        projects_df = pd.DataFrame(project_service.list_projects(limit=500))
        assignments_df = pd.DataFrame(project_service.list_assignments(limit=500))
        project_filter_df = projects_df.copy()
        if not project_filter_df.empty:
            project_filter_df["client_name"] = project_filter_df["client_name"].fillna("Unspecified")

        render_panel_intro(
            "Project Overview",
            "See the live status of every project, current delivery mix, and the full project directory in one place.",
            eyebrow="Overview",
        )
        _render_selection_spacer()
        overview_filter_col1, overview_filter_col2, overview_filter_col3 = st.columns(3)
        with overview_filter_col1:
            selected_statuses = st.multiselect(
                "Status filter",
                PROJECT_STATUSES,
                key=PROJECT_OVERVIEW_STATUS_FILTER_KEY,
                placeholder="All statuses",
            )
        with overview_filter_col2:
            selected_clients = st.multiselect(
                "Client filter",
                sorted(project_filter_df["client_name"].dropna().unique().tolist()) if not project_filter_df.empty else [],
                key=PROJECT_OVERVIEW_CLIENT_FILTER_KEY,
                placeholder="All clients",
            )
        with overview_filter_col3:
            selected_project_types = st.multiselect(
                "Project type filter",
                sorted(project_filter_df["project_type"].dropna().unique().tolist()) if not project_filter_df.empty else [],
                key=PROJECT_OVERVIEW_TYPE_FILTER_KEY,
                placeholder="All project types",
            )

        filtered_projects_df = _filter_projects_overview(
            project_filter_df,
            statuses=selected_statuses,
            clients=selected_clients,
            project_types=selected_project_types,
        )
        filtered_project_ids = set(filtered_projects_df["project_id"].tolist()) if not filtered_projects_df.empty else set()
        filtered_assignments_df = (
            assignments_df[assignments_df["project_id"].isin(filtered_project_ids)].reset_index(drop=True)
            if not assignments_df.empty and filtered_project_ids
            else pd.DataFrame(columns=assignments_df.columns if not assignments_df.empty else [])
        )
        overview_subtitle = _filters_summary_text(selected_statuses, selected_clients, selected_project_types)
        overview_summary_items = _project_stat_items(filtered_projects_df)
        status_overview_frame = _project_status_overview_frame(filtered_projects_df)
        type_overview_frame = _project_type_overview_frame(filtered_projects_df)
        client_overview_frame = _project_client_overview_frame(filtered_projects_df)

        render_stat_band(_project_stat_items(filtered_projects_df))

        with st.container(key="project_overview_export_shell"):
            render_panel_intro(
                "Export Reports",
                "Download the current filtered overview as Word, Excel, CSV, or PDF.",
                eyebrow="Export",
                class_name="panel-intro--section-break",
            )
            _render_selection_spacer()
            export_sections = [
                ("Status Summary", status_overview_frame),
                ("Project Types", type_overview_frame),
                ("Client Portfolio", client_overview_frame),
                ("Project Directory", filtered_projects_df),
                ("Assignment Directory", filtered_assignments_df),
            ]
            workbook_sheets = {
                "Project Overview": filtered_projects_df,
                "Assignments": filtered_assignments_df,
                "Status Summary": status_overview_frame,
                "Project Types": type_overview_frame,
                "Client Portfolio": client_overview_frame,
            }
            word_bytes = _build_word_report_bytes(
                "Project Overview Report",
                overview_subtitle,
                overview_summary_items,
                export_sections,
            )
            pdf_bytes = _build_pdf_report_bytes(
                "Project Overview Report",
                overview_subtitle,
                overview_summary_items,
                export_sections,
            )
            xlsx_bytes = _frames_to_xlsx_bytes(workbook_sheets)

            export_col1, export_col2, export_col3 = st.columns(3)
            with export_col1:
                st.download_button(
                    "Word report",
                    data=word_bytes,
                    file_name="project_overview_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="project_overview_word_export",
                    width="stretch",
                )
            with export_col2:
                st.download_button(
                    "Excel workbook",
                    data=xlsx_bytes,
                    file_name="project_overview_report.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="project_overview_excel_export",
                    width="stretch",
                )
            with export_col3:
                st.download_button(
                    "PDF report",
                    data=pdf_bytes,
                    file_name="project_overview_report.pdf",
                    mime="application/pdf",
                    key="project_overview_pdf_export",
                    width="stretch",
                )

            csv_col1, csv_col2 = st.columns(2)
            with csv_col1:
                st.download_button(
                    "Projects CSV",
                    data=_dataframe_to_csv_bytes(filtered_projects_df),
                    file_name="project_directory.csv",
                    mime="text/csv",
                    key="project_overview_projects_csv_export",
                    width="stretch",
                )
            with csv_col2:
                st.download_button(
                    "Assignments CSV",
                    data=_dataframe_to_csv_bytes(filtered_assignments_df),
                    file_name="project_assignments.csv",
                    mime="text/csv",
                    key="project_overview_assignments_csv_export",
                    width="stretch",
                )

        if filtered_projects_df.empty:
            _render_selection_spacer()
            st.caption("No projects match the current filters.")
        else:
            with st.container(key="project_overview_status_shell"):
                render_panel_intro(
                    "Status Summary",
                    "Track how many projects are active, planned, on hold, or closed, together with assignment volume.",
                    eyebrow="Status",
                    class_name="panel-intro--section-break",
                )
                _render_selection_spacer()
                render_table(status_overview_frame, max_visible_rows=6)

            overview_col1, overview_col2 = st.columns(2)
            with overview_col1:
                render_panel_intro(
                    "Project Types",
                    "Understand where delivery volume is concentrated by project type.",
                    eyebrow="Mix",
                    class_name="panel-intro--section-break",
                )
                _render_selection_spacer()
                render_table(type_overview_frame, max_visible_rows=6)
            with overview_col2:
                render_panel_intro(
                    "Client Portfolio",
                    "See which clients hold the largest number of active and planned projects.",
                    eyebrow="Clients",
                    class_name="panel-intro--section-break",
                )
                _render_selection_spacer()
                render_table(client_overview_frame, max_visible_rows=6)

            render_panel_intro(
                "Project Directory",
                "Review all recorded project fields, including status, phase, timeline, and assignment load.",
                eyebrow="Directory",
                class_name="panel-intro--section-break",
            )
            _render_selection_spacer()
            project_overview_cols = [
                "project_code",
                "project_name",
                "project_short_name",
                "client_name",
                "implementing_partner",
                "project_type",
                "status",
                "phase_number",
                "assignment_count",
                "start_date",
                "end_date",
            ]
            render_table(apply_text_filter(filtered_projects_df[project_overview_cols], "Search project overview"))

            if not filtered_assignments_df.empty:
                render_panel_intro(
                    "Assignment Directory",
                    "Review all current and historical project assignments from the same workspace.",
                    eyebrow="Assignments",
                    class_name="panel-intro--section-break",
                )
                _render_selection_spacer()
                assignment_overview_cols = [
                    "project_code",
                    "project_name",
                    "surveyor_code",
                    "surveyor_name",
                    "role",
                    "work_province_name",
                    "status",
                    "start_date",
                    "end_date",
                ]
                render_table(apply_text_filter(filtered_assignments_df[assignment_overview_cols], "Search project assignments"))

    elif active_view == "New Project":
        render_panel_intro("Create New Project", eyebrow=None)
        with st.form(PROJECT_FORM_KEY, clear_on_submit=False):
            vf.render_form_error_summary(PROJECT_FORM_KEY)
            col1, col2 = st.columns(2)
            with col1:
                project_name = vf.text_input(PROJECT_FORM_KEY, PROJECT_NAME_KEY, "Project name", required=True)
                project_short_name = vf.text_input(PROJECT_FORM_KEY, PROJECT_SHORT_NAME_KEY, "Project short name")
                client_name = vf.text_input(PROJECT_FORM_KEY, PROJECT_CLIENT_KEY, "Client name")
                implementing_partner = vf.text_input(PROJECT_FORM_KEY, PROJECT_PARTNER_KEY, "Implementing partner")
                project_type = vf.selectbox(PROJECT_FORM_KEY, PROJECT_TYPE_KEY, "Project type", PROJECT_TYPES, required=True)
            with col2:
                status = vf.selectbox(PROJECT_FORM_KEY, PROJECT_STATUS_KEY, "Status", PROJECT_STATUSES, required=True)
                start_date = vf.date_input(PROJECT_FORM_KEY, PROJECT_START_KEY, "Start", value=None)
                end_date = vf.date_input(PROJECT_FORM_KEY, PROJECT_END_KEY, "End", value=None)
                project_document_link = vf.text_input(PROJECT_FORM_KEY, PROJECT_DOCUMENT_KEY, "Document link")
                notes = vf.text_area(PROJECT_FORM_KEY, PROJECT_NOTES_KEY, "Notes", height=130)
            submitted = st.form_submit_button("Create project", width="stretch")
        if submitted:
            payload = {
                "project_name": project_name.strip(),
                "project_short_name": project_short_name.strip() or None,
                "client_name": client_name.strip() or None,
                "implementing_partner": implementing_partner.strip() or None,
                "project_type": project_type,
                "start_date": start_date,
                "end_date": end_date,
                "status": status,
                "notes": notes.strip() or None,
                "project_document_link": project_document_link.strip() or None,
            }
            errors = vf.required_errors({PROJECT_NAME_KEY: (payload["project_name"], "Enter the project name before creating it.")})
            if payload["start_date"] and payload["end_date"] and payload["start_date"] > payload["end_date"]:
                errors[PROJECT_END_KEY] = "Choose an end date that is the same as or after the start date."
            if errors:
                errors[vf.FORM_MESSAGE_KEY] = "Please fix the highlighted project fields."
                vf.stop_with_form_errors(PROJECT_FORM_KEY, errors)
            try:
                created = project_service.create_project(get_current_user(), payload)
            except Exception as exc:
                message = vf.user_friendly_error_message(
                    exc,
                    "We could not create this project right now. Check the project details and try again.",
                )
                vf.stop_with_form_errors(
                    PROJECT_FORM_KEY,
                    vf.field_errors_from_message(
                        message,
                        {
                            PROJECT_NAME_KEY: ("project", "project name"),
                            PROJECT_SHORT_NAME_KEY: ("short name", "generated code"),
                            PROJECT_CLIENT_KEY: ("client",),
                            PROJECT_START_KEY: ("start", "year"),
                        },
                    ),
                )
            vf.clear_form_errors(PROJECT_FORM_KEY)
            st.success(f"Project created with code {created['project_code']}.")
            st.rerun()

    elif active_view == "Assignments":
        render_panel_intro("Smart Assignment Builder", eyebrow="Create")
        provinces = ProvinceRepository().list_all()
        province_map = {f"{item['province_name']} ({item['province_code']})": item["province_code"] for item in provinces}
        province_label_by_code = {item["province_code"]: f"{item['province_name']} ({item['province_code']})" for item in provinces}
        projects = project_service.list_projects(limit=1000)
        project_map = {f"{item['project_name']} ({item['project_code']})": item["project_id"] for item in projects}
        if not project_map:
            st.warning("Create a project first, then come back to assign surveyors.")
            return
        if not province_map:
            st.warning("Add province records first, then use the assignment builder.")
            return

        with st.container(key="assignment_scope_shell"):
            selector_col1, selector_col2 = st.columns(2)
            with selector_col1:
                project_label = st.selectbox(
                    "Project",
                    list(project_map.keys()),
                    key=ASSIGNMENT_PROJECT_KEY,
                )
            with selector_col2:
                province_options = [(item["province_code"], province_label_by_code[item["province_code"]]) for item in provinces]
                selected_province_codes = _render_checkbox_selector(
                    "Work provinces",
                    province_options,
                    state_key=ASSIGNMENT_PROVINCES_KEY,
                    control_prefix="assignment_province",
                    help_text="Keep this list open while you finish choosing the provinces you want to work with.",
                    columns=1 if len(province_options) <= 8 else 2,
                )

        selected_project_id = project_map.get(project_label)
        selected_province_labels = [province_label_by_code[code] for code in selected_province_codes if code in province_label_by_code]
        candidate_rows = SurveyorService().list_assignment_candidates(selected_province_codes, limit=1000) if selected_province_codes else []
        candidate_by_id = {item["surveyor_id"]: item for item in candidate_rows}
        grouped_candidates = _group_assignment_candidates(candidate_rows)
        province_caption = ", ".join(selected_province_labels[:4]) + (" ..." if len(selected_province_labels) > 4 else "")
        st.info(
            f"{len(candidate_rows)} surveyor(s) available across {len(selected_province_codes)} province(s)"
            + (f": {province_caption}." if province_caption else ".")
            + " Each province now has its own clean selection block, so the list stays stable while you complete your choices."
        )

        if candidate_rows:
            with st.container(key="assignment_coverage_shell"):
                render_panel_intro(
                    "Province Coverage",
                    "A clean summary of how many surveyors are available in each selected province.",
                    eyebrow="Overview",
                    class_name="panel-intro--section-break",
                )
                _render_selection_spacer()
                province_summary_df = _province_summary_frame(candidate_rows)
                if not province_summary_df.empty:
                    render_table(province_summary_df, max_visible_rows=6)
        else:
            st.caption("No surveyors are currently mapped to the selected provinces.")

        with st.container(key="assignment_selection_shell"):
            render_panel_intro(
                "Surveyor Selection",
                "Choose surveyors inside each province block, then confirm the shared assignment details below.",
                eyebrow="Selection",
                class_name="panel-intro--section-break",
            )
            _render_selection_spacer()

            selected_surveyor_ids_state = [
                surveyor_id
                for surveyor_id in st.session_state.get(ASSIGNMENT_SURVEYORS_KEY, [])
                if surveyor_id in candidate_by_id
            ]
            st.session_state[ASSIGNMENT_SURVEYORS_KEY] = selected_surveyor_ids_state

            if grouped_candidates:
                st.caption("Use the province-level actions when you want to mark or clear a whole block quickly.")
                selected_surveyor_ids: list[int] = []
                for province_code, province_name, province_rows in grouped_candidates:
                    province_surveyor_ids = {int(item["surveyor_id"]) for item in province_rows}
                    province_selected_ids = [
                        str(surveyor_id)
                        for surveyor_id in selected_surveyor_ids_state
                        if surveyor_id in province_surveyor_ids
                    ]
                    ready_docs = sum(1 for item in province_rows if int(item.get("document_count", 0)) >= 4)
                    active_accounts = sum(int(item.get("active_account_count", 0)) for item in province_rows)
                    active_projects = sum(int(item.get("active_project_count", 0)) for item in province_rows)
                    block_label = f"{province_name} ({province_code})" if province_code != "NA" else province_name
                    group_options = [(str(item["surveyor_id"]), _candidate_selection_label(item)) for item in province_rows]
                    control_prefix = f"assignment_surveyor_{province_code}"
                    with st.expander(
                        f"{block_label} - {len(province_rows)} surveyor(s)",
                        expanded=len(grouped_candidates) <= 2 or bool(province_selected_ids),
                    ):
                        st.caption(
                            f"Ready docs: {ready_docs}   Active accounts: {active_accounts}   Active projects: {active_projects}"
                        )
                        action_col1, action_col2, action_col3 = st.columns([0.8, 0.8, 2.2])
                        with action_col1:
                            if st.button("Select all", key=f"{control_prefix}_select_all", width="stretch"):
                                _set_checkbox_grid_values(group_options, control_prefix=control_prefix, checked=True)
                        with action_col2:
                            if st.button("Clear all", key=f"{control_prefix}_clear_all", width="stretch"):
                                _set_checkbox_grid_values(group_options, control_prefix=control_prefix, checked=False)
                        with action_col3:
                            st.caption("These actions only affect this province block.")
                        search_term = st.text_input(
                            "Search in this province",
                            key=f"{control_prefix}_search",
                            placeholder="Search by surveyor name or code",
                        )
                        filtered_rows = _filter_assignment_rows(province_rows, search_term)
                        filtered_group_options = [
                            (str(item["surveyor_id"]), _candidate_selection_label(item))
                            for item in filtered_rows
                        ]
                        if search_term.strip():
                            st.caption(f"Showing {len(filtered_rows)} of {len(province_rows)} surveyor(s) in this province.")
                        _render_checkbox_grid(
                            filtered_group_options,
                            control_prefix=control_prefix,
                            selected_values=province_selected_ids,
                            columns=1 if len(filtered_rows) <= 6 else 2,
                        )
                        if not filtered_rows:
                            st.caption("No surveyors match this search.")
                        selected_group_ids = _read_checkbox_grid_values(group_options, control_prefix=control_prefix)
                        selected_surveyor_ids.extend(int(surveyor_id) for surveyor_id in selected_group_ids)
                selected_surveyor_ids = list(dict.fromkeys(selected_surveyor_ids))
                st.session_state[ASSIGNMENT_SURVEYORS_KEY] = selected_surveyor_ids
                st.caption(f"{len(selected_surveyor_ids)} surveyor(s) currently selected.")
            else:
                selected_surveyor_ids = []
                st.session_state[ASSIGNMENT_SURVEYORS_KEY] = []
                st.caption("Select one or more provinces first. The grouped surveyor blocks will appear here.")

        selected_ready_docs = sum(
            1
            for surveyor_id in selected_surveyor_ids
            if int(candidate_by_id.get(surveyor_id, {}).get("document_count", 0)) >= 4
        )
        selected_active_accounts = sum(
            int(candidate_by_id.get(surveyor_id, {}).get("active_account_count", 0))
            for surveyor_id in selected_surveyor_ids
        )
        with st.container(key="assignment_snapshot_shell"):
            render_panel_intro(
                "Assignment Snapshot",
                "A live summary of the current province and surveyor selection before you save.",
                eyebrow="Live",
                class_name="panel-intro--section-break",
            )
            _render_selection_spacer()
            render_stat_band(
                [
                    ("Selected provinces", f"{len(selected_province_codes)}"),
                    ("Selected surveyors", f"{len(selected_surveyor_ids)}"),
                    ("Ready documents", f"{selected_ready_docs}"),
                    ("Assignments to create", f"{len(selected_surveyor_ids)}"),
                    ("Active accounts", f"{selected_active_accounts}"),
                ]
            )
            _render_selection_spacer()

        with st.container(key="assignment_details_shell"):
            render_panel_intro(
                "Assignment Details",
                "These details apply to all currently selected surveyors in one save.",
                eyebrow="Details",
                class_name="panel-intro--section-break",
            )
            _render_selection_spacer()
            vf.render_form_error_summary(ASSIGNMENT_FORM_KEY)
            col1, col2 = st.columns(2)
            with col1:
                role = vf.text_input(ASSIGNMENT_FORM_KEY, ASSIGNMENT_ROLE_KEY, "Assignment role", required=True, value="Surveyor")
                status = vf.selectbox(ASSIGNMENT_FORM_KEY, ASSIGNMENT_STATUS_KEY, "Assignment status", ASSIGNMENT_STATUSES, required=True)
            with col2:
                extra_province_codes = _render_checkbox_selector(
                    "Additional provinces beyond the selected set",
                    [
                        (item["province_code"], province_label_by_code[item["province_code"]])
                        for item in provinces
                        if item["province_code"] not in selected_province_codes
                    ],
                    state_key=ASSIGNMENT_EXTRA_PROVINCES_KEY,
                    control_prefix="assignment_extra_province",
                    help_text="Use this only when the same assignment should also cover more provinces.",
                    columns=1 if len(provinces) <= 8 else 2,
                )
                start_date = vf.date_input(ASSIGNMENT_FORM_KEY, ASSIGNMENT_START_KEY, "Start", value=None, required=True)
                end_date = vf.date_input(ASSIGNMENT_FORM_KEY, ASSIGNMENT_END_KEY, "End", value=None)
                notes = vf.text_area(ASSIGNMENT_FORM_KEY, ASSIGNMENT_NOTES_KEY, "Assignment notes", height=120)

            conflict_rows = project_service.list_assignment_conflicts(
                selected_project_id,
                selected_surveyor_ids,
                start_date,
                end_date,
                limit=500,
            ) if selected_surveyor_ids and start_date else []

            if conflict_rows:
                same_project_count = len({item["surveyor_id"] for item in conflict_rows if item.get("same_project")})
                overlap_count = len({item["surveyor_id"] for item in conflict_rows if item.get("overlaps_window")})
                with st.container(key="assignment_conflict_shell"):
                    render_panel_intro(
                        "Conflict Review",
                        "Some selected surveyors already have an assignment on this project or overlap the chosen date window.",
                        eyebrow="Review",
                        class_name="panel-intro--section-break",
                    )
                    _render_selection_spacer()
                    render_stat_band(
                        [
                            ("Blocked surveyors", f"{len({item['surveyor_id'] for item in conflict_rows})}"),
                            ("Same project", f"{same_project_count}"),
                            ("Date overlap", f"{overlap_count}"),
                        ]
                    )
                    _render_selection_spacer("0.42rem")
                    render_table(_assignment_conflict_frame(conflict_rows), max_visible_rows=8)
                    st.warning("Resolve the highlighted conflicts before saving this batch.")
            elif selected_surveyor_ids and start_date:
                st.caption("No duplicate project assignments or overlapping date conflicts were found for the current selection.")

            submitted = st.button("Assign selected surveyors", key="assignment_submit", width="stretch")
        if submitted:
            st.session_state[ASSIGNMENT_SURVEYORS_KEY] = selected_surveyor_ids
            errors = vf.required_errors(
                {
                    ASSIGNMENT_PROVINCES_KEY: (selected_province_codes, "Select at least one province before assigning surveyors."),
                    ASSIGNMENT_SURVEYORS_KEY: (selected_surveyor_ids, "Select at least one surveyor to assign to this project."),
                    ASSIGNMENT_START_KEY: (start_date, "Choose the assignment start date before saving it."),
                }
            )
            if not candidate_rows:
                errors[ASSIGNMENT_SURVEYORS_KEY] = "No surveyors are available in the selected provinces yet."
            if errors:
                errors[vf.FORM_MESSAGE_KEY] = "Please fix the highlighted assignment fields."
                vf.stop_with_form_errors(ASSIGNMENT_FORM_KEY, errors)
            if end_date and start_date and end_date < start_date:
                vf.stop_with_form_errors(
                    ASSIGNMENT_FORM_KEY,
                    {
                        vf.FORM_MESSAGE_KEY: "Please fix the assignment date range.",
                        ASSIGNMENT_END_KEY: "Choose an end date that is the same as or after the start date.",
                    },
                )
            if conflict_rows:
                vf.stop_with_form_errors(
                    ASSIGNMENT_FORM_KEY,
                    {
                        vf.FORM_MESSAGE_KEY: "Resolve assignment conflicts before saving this batch.",
                        ASSIGNMENT_SURVEYORS_KEY: "One or more selected surveyors already belong to this project or overlap the chosen date window.",
                    },
                )
            selected_surveyor_ids = list(dict.fromkeys(selected_surveyor_ids))
            payloads = []
            for surveyor_id in selected_surveyor_ids:
                candidate = candidate_by_id.get(surveyor_id)
                if not candidate:
                    continue
                primary_province_code = candidate.get("availability_province_code")
                if not primary_province_code:
                    continue
                extra_codes = [
                    code
                    for code in dict.fromkeys(selected_province_codes + extra_province_codes)
                    if code != primary_province_code
                ]
                payloads.append(
                    {
                        "project_id": selected_project_id,
                        "surveyor_id": surveyor_id,
                        "role": role.strip() or "Surveyor",
                        "work_province_code": primary_province_code,
                        "extra_province_codes": extra_codes,
                        "start_date": start_date,
                        "end_date": end_date,
                        "status": status,
                        "notes": notes.strip() or None,
                    }
                )
            if not payloads:
                vf.stop_with_form_errors(
                    ASSIGNMENT_FORM_KEY,
                    {
                        vf.FORM_MESSAGE_KEY: "No valid surveyors are ready to save from the current selection.",
                        ASSIGNMENT_SURVEYORS_KEY: "Tick one or more surveyors with a valid province assignment.",
                    },
                )
            try:
                created_assignments = project_service.create_assignments(get_current_user(), payloads)
            except Exception as exc:
                message = vf.user_friendly_error_message(
                    exc,
                    "We could not save these assignments right now. Check the selected project, surveyors, and province.",
                )
                vf.stop_with_form_errors(
                    ASSIGNMENT_FORM_KEY,
                    vf.field_errors_from_message(
                        message,
                        {
                            ASSIGNMENT_PROJECT_KEY: ("project",),
                            ASSIGNMENT_SURVEYORS_KEY: ("surveyor",),
                            ASSIGNMENT_PROVINCES_KEY: ("province",),
                            ASSIGNMENT_END_KEY: ("date",),
                        },
                    ),
                )
            vf.clear_form_errors(ASSIGNMENT_FORM_KEY)
            st.session_state[ASSIGNMENT_SURVEYORS_KEY] = []
            st.session_state[ASSIGNMENT_EXTRA_PROVINCES_KEY] = []
            st.success(f"{len(created_assignments)} surveyor assignment(s) created for {project_label}.")
            st.rerun()

    else:
        projects_df = pd.DataFrame(project_service.list_projects(limit=500))
        assignments_df = pd.DataFrame(project_service.list_assignments(limit=500))
        render_panel_intro("Projects", eyebrow="Directory")
        if not projects_df.empty:
            project_cols = [
                "project_code",
                "project_name",
                "project_short_name",
                "client_name",
                "project_type",
                "status",
                "phase_number",
                "assignment_count",
                "start_date",
                "end_date",
            ]
            render_table(apply_text_filter(projects_df[project_cols], "Search projects"))
        else:
            render_table(projects_df)

        render_panel_intro("Assignments", eyebrow="Directory", class_name="panel-intro--section-break")
        if not assignments_df.empty:
            assignment_cols = [
                "project_code",
                "project_name",
                "surveyor_code",
                "surveyor_name",
                "role",
                "work_province_name",
                "extra_province_codes",
                "status",
                "start_date",
                "end_date",
            ]
            render_table(apply_text_filter(assignments_df[assignment_cols], "Search assignments"))
        else:
            render_table(assignments_df)
